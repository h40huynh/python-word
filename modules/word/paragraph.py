from __future__ import annotations
import re

from modules.word.run import Run
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


class Paragraph:
    def __init__(self, paragraph) -> None:
        self.__paragraph = paragraph
        self.text = paragraph.text

    def __getitem__(self, item) -> Run:
        return Run(self.__paragraph.runs[item])

    def set_style(self, style: str) -> Paragraph:
        self.__paragraph.style = style
        return self

    def add_run(self, text: str, is_image: bool = False) -> Run:
        return Run(self.__paragraph.add_run(text))

    def add_picture(self, filename: str, width: float = None, height: float = None) -> Run:
        if height:
            self.__paragraph.add_run().add_picture(
                filename, width=Inches(width), height=Inches(height))
        else:
            self.__paragraph.add_run().add_picture(filename, width=Inches(width))
        return self

    def set_alignment(self, alignment: str) -> Paragraph:
        if alignment == "left":
            self.__paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "right":
            self.__paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "center":
            self.__paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "justify":
            self.__paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self

    def replace_regex(self, regex: re.Pattern, replacement: str) -> Paragraph:
        while True:
            text = self.__paragraph.text
            match = regex.search(text)
            if not match:
                break

            runs = iter(self.__paragraph.runs)
            start, end = match.start(), match.end()

            for run in runs:
                run_len = len(run.text)
                if start < run_len:
                    break
                start, end = start - run_len, end - run_len

            run_text = run.text
            run_len = len(run_text)
            run.text = "%s%s%s" % (
                run_text[:start], replacement, run_text[end:])
            end -= run_len

            for run in runs:
                if end <= 0:
                    break
                run_text = run.text
                run_len = len(run_text)
                run.text = run_text[end:]
                end -= run_len
        return self

    def replace_str(self, find: str, replacement: str) -> Paragraph:
        return self.replace_regex(re.compile(find), replacement)
