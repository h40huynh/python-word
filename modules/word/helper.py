from typing import Any
from docx import Document
from docx.shared import Inches
from modules.word.paragraph import Paragraph
from modules.word.table import Table


class WordHelper:
    def __init__(self, document_path: str) -> None:
        with open(document_path, "rb") as f:
            self.__document = Document(f)
        self.paragraphs = [Paragraph(p) for p in self.__document.paragraphs]
        self.tables = [Table(t) for t in self.__document.tables]

    def add_paragraph(self, text: str) -> Paragraph:
        return Paragraph(self.__document.add_paragraph(text))

    def add_heading(self, text: str, level: int) -> Paragraph:
        return Paragraph(self.__document.add_heading(text, level))

    def add_picture(self, filename: str, width: float = None, height: float = None) -> Any:
        if height:
            return (self.__document.add_picture(filename, width=Inches(width), height=Inches(height)))
        else:
            return self.__document.add_picture(filename, width=Inches(width))

    def add_page_break(self) -> None:
        self.__document.add_page_break()

    def save(self, filename: str) -> None:
        self.__document.save(filename)
